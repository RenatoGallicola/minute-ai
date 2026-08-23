"""
export.py
---------
Exports transcript and summary to md, txt, docx, pdf and srt formats.
"""

import os
import re
from datetime import datetime
from xml.sax.saxutils import escape as xml_escape

from src.errors import ExportError
from src.logger import get_logger
from src.naming import output_stem, unique_path
from src.transcribe import format_duration


FORMATS = ("md", "txt", "docx", "pdf", "srt")
EXTENSIONS = {fmt: f".{fmt}" for fmt in FORMATS}


def formats_for(fmt: str) -> list[str]:
    """Expands 'all' into every concrete format, or wraps a single one."""
    return list(FORMATS) if fmt == "all" else [fmt]


def export(
    audio_path: str,
    meeting_name: str,
    transcript: str,
    summary: str,
    language: str,
    output_dir: str,
    fmt: str,
    export_content: str,
    duration_seconds: float = 0.0,
    segments: list = None,
) -> list[str]:
    """
    Exports results to the specified formats.

    Args:
        audio_path:       Path to the original audio file
        meeting_name:     Human-readable meeting name
        transcript:       Full transcript
        summary:          Structured summary (may be empty)
        language:         Detected language
        output_dir:       Output folder
        fmt:              Format: 'md', 'txt', 'docx', 'pdf', 'srt', 'all'
        export_content:   Content to include: 'full' or 'summary'
        duration_seconds: Audio duration, shown in the file header
        segments:         Timestamped whisperX segments (required for 'srt')

    Returns:
        List of paths of created files
    """
    log = get_logger()

    try:
        os.makedirs(output_dir, exist_ok=True)
    except OSError as exc:
        raise ExportError(f"Cannot create the output folder '{output_dir}': {exc}") from exc

    audio_stem = os.path.splitext(os.path.basename(audio_path))[0]
    base_filename = output_stem(meeting_name)

    # Determine what content to include
    include_transcript = export_content == "full"
    include_summary = bool(summary)

    if not include_summary and not include_transcript:
        log.warning(
            "Nothing to export: 'summary only' was requested but no summary was produced. "
            "Writing the header only — check that Ollama is running."
        )

    meta = _Meta(meeting_name, audio_stem, language, duration_seconds)
    writers = {
        "md": lambda path: _write_markdown(path, meta, summary, transcript, include_transcript, include_summary),
        "txt": lambda path: _write_txt(path, meta, summary, transcript, include_transcript, include_summary),
        "docx": lambda path: _write_docx(path, meta, summary, transcript, include_transcript, include_summary),
        "pdf": lambda path: _write_pdf(path, meta, summary, transcript, include_transcript, include_summary),
        "srt": lambda path: _write_srt(path, segments),
    }

    created_files = []
    for name in formats_for(fmt):
        if name == "srt" and not segments:
            log.warning("Skipping .srt export: no timestamped segments are available.")
            continue

        # unique_path keeps two sources that resolve to the same name in the
        # same minute (meeting.mp3 / meeting.wav) from overwriting each other.
        path = unique_path(os.path.join(output_dir, base_filename + EXTENSIONS[name]))
        try:
            written = writers[name](path)
        except ExportError:
            raise
        except OSError as exc:
            raise ExportError(f"Cannot write '{path}': {exc}") from exc

        if written is False:  # optional dependency missing — already logged
            continue
        created_files.append(path)
        log.info(f"      Exported: {path}")

    return created_files


class _Meta:
    """The header fields every format shares."""

    def __init__(self, meeting_name, audio_stem, language, duration_seconds):
        self.meeting_name = meeting_name
        self.audio_stem = audio_stem
        self.language = language
        self.duration_seconds = duration_seconds
        self.date_str = datetime.now().strftime("%Y-%m-%d %H:%M")

    @property
    def line(self) -> str:
        parts = [f"Date: {self.date_str}", f"Language: {self.language}", f"Source: {self.audio_stem}"]
        if self.duration_seconds:
            parts.insert(1, f"Duration: {format_duration(self.duration_seconds)}")
        return "  |  ".join(parts)


# ─────────────────────────────────────────────
# Markdown
# ─────────────────────────────────────────────

def _write_markdown(path, meta, summary, transcript, include_transcript, include_summary):
    content = f"# {meta.meeting_name}\n\n"
    content += f"> {meta.line}\n\n---\n\n"

    if include_summary and summary:
        content += summary + "\n\n---\n\n"

    if include_transcript:
        content += "## Full Transcript\n\n" + transcript + "\n"

    with open(path, "w", encoding="utf-8") as f:
        f.write(content)


# ─────────────────────────────────────────────
# Plain text
# ─────────────────────────────────────────────

def _write_txt(path, meta, summary, transcript, include_transcript, include_summary):
    lines = [f"MEETING: {meta.meeting_name}", meta.line, "=" * 60, ""]

    if include_summary and summary:
        clean = summary.replace("## ", "").replace("**", "")
        lines += ["SUMMARY", "-" * 40, clean, "", "=" * 60, ""]

    if include_transcript:
        lines += ["TRANSCRIPT", "-" * 40, transcript]

    with open(path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))


# ─────────────────────────────────────────────
# SRT subtitles
# ─────────────────────────────────────────────

def _write_srt(path, segments):
    """Writes timestamped subtitles from the aligned whisperX segments."""
    blocks = []
    index = 1
    for seg in segments or []:
        text = (seg.get("text") or "").strip()
        if not text:
            continue
        speaker = seg.get("speaker")
        label = f"[{speaker}] " if speaker else ""
        start = _srt_timestamp(seg.get("start", 0.0))
        end = _srt_timestamp(seg.get("end", seg.get("start", 0.0)))
        blocks.append(f"{index}\n{start} --> {end}\n{label}{text}\n")
        index += 1

    with open(path, "w", encoding="utf-8") as f:
        f.write("\n".join(blocks))


def _srt_timestamp(seconds: float) -> str:
    total_ms = max(int(round(float(seconds) * 1000)), 0)
    hours, rest = divmod(total_ms, 3_600_000)
    minutes, rest = divmod(rest, 60_000)
    secs, millis = divmod(rest, 1000)
    return f"{hours:02d}:{minutes:02d}:{secs:02d},{millis:03d}"


# ─────────────────────────────────────────────
# DOCX
# ─────────────────────────────────────────────

def _write_docx(path, meta, summary, transcript, include_transcript, include_summary):
    log = get_logger()
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        log.error("python-docx is not installed. Run: pip install python-docx")
        return False

    doc = Document()

    # Title
    title = doc.add_heading(meta.meeting_name, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Metadata
    doc.add_paragraph().add_run(meta.line).italic = True
    doc.add_paragraph()  # spacer

    # Summary
    if include_summary and summary:
        _docx_add_markdown(doc, summary)
        if include_transcript:
            doc.add_page_break()

    # Transcript
    if include_transcript:
        doc.add_heading("Full Transcript", level=1)
        for line in transcript.split("\n\n"):
            if line.strip():
                p = doc.add_paragraph()
                speaker, text = _split_speaker(line)
                if speaker is not None:
                    p.add_run(speaker + ": ").bold = True
                    p.add_run(text)
                else:
                    p.add_run(line.strip())

    doc.save(path)


def _docx_add_markdown(doc, text):
    """Converts basic Markdown to python-docx elements."""
    for line in text.split("\n"):
        line = line.rstrip()

        if line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.lstrip().startswith(("- ", "* ")):
            p = doc.add_paragraph(style="List Bullet")
            _docx_add_inline(p, line.lstrip()[2:])
        elif line == "---":
            doc.add_paragraph("─" * 40)
        elif line.strip() == "":
            doc.add_paragraph()
        else:
            p = doc.add_paragraph()
            _docx_add_inline(p, line)


def _docx_add_inline(paragraph, text):
    """Handles **bold** inline markdown in a paragraph."""
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            paragraph.add_run(part[2:-2]).bold = True
        else:
            paragraph.add_run(part)


# ─────────────────────────────────────────────
# PDF
# ─────────────────────────────────────────────

def _write_pdf(path, meta, summary, transcript, include_transcript, include_summary):
    log = get_logger()
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import cm
        from reportlab.lib import colors
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable, PageBreak
    except ImportError:
        log.error("reportlab is not installed. Run: pip install reportlab")
        return False

    doc = SimpleDocTemplate(
        path,
        pagesize=A4,
        rightMargin=2*cm, leftMargin=2*cm,
        topMargin=2*cm, bottomMargin=2*cm,
        title=meta.meeting_name,
    )

    styles = getSampleStyleSheet()
    style_title = ParagraphStyle("MeetingTitle", parent=styles["Title"], fontSize=18, spaceAfter=6)
    style_meta = ParagraphStyle("Meta", parent=styles["Normal"], fontSize=9, textColor=colors.grey, spaceAfter=12)
    style_h1 = ParagraphStyle("H1", parent=styles["Heading1"], fontSize=14, spaceBefore=12, spaceAfter=6)
    style_h2 = ParagraphStyle("H2", parent=styles["Heading2"], fontSize=12, spaceBefore=10, spaceAfter=4)
    style_body = ParagraphStyle("Body", parent=styles["Normal"], fontSize=10, spaceAfter=6, leading=14)

    story = []

    # Title
    story.append(Paragraph(_pdf_text(meta.meeting_name), style_title))
    story.append(Paragraph(_pdf_text(meta.line).replace("  |  ", " &nbsp;|&nbsp; "), style_meta))
    story.append(HRFlowable(width="100%", thickness=1, color=colors.lightgrey))
    story.append(Spacer(1, 12))

    # Summary
    if include_summary and summary:
        for line in summary.split("\n"):
            line = line.rstrip()
            if line.startswith("## "):
                story.append(Paragraph(_pdf_text(line[3:]), style_h2))
            elif line.startswith("# "):
                story.append(Paragraph(_pdf_text(line[2:]), style_h1))
            elif line.lstrip().startswith(("- ", "* ")):
                story.append(Paragraph(f"• {_pdf_inline(line.lstrip()[2:])}", style_body))
            elif line == "---":
                story.append(HRFlowable(width="100%", thickness=0.5, color=colors.lightgrey))
            elif line.strip():
                story.append(Paragraph(_pdf_inline(line), style_body))
            else:
                story.append(Spacer(1, 6))

        if include_transcript:
            story.append(PageBreak())

    # Transcript
    if include_transcript:
        story.append(Paragraph("Full Transcript", style_h1))
        story.append(HRFlowable(width="100%", thickness=0.5, color=colors.lightgrey))
        story.append(Spacer(1, 8))

        for block in transcript.split("\n\n"):
            if not block.strip():
                continue
            speaker, text = _split_speaker(block)
            if speaker is not None:
                content = f"<b>{_pdf_text(speaker)}:</b> {_pdf_text(text)}"
            else:
                content = _pdf_text(block.strip())
            story.append(Paragraph(content, style_body))

    doc.build(story)


def _pdf_text(text: str) -> str:
    """Escapes text for reportlab's mini-HTML parser.

    Paragraph() parses its input as markup, so a bare '&' (as in 'R&D') or a
    '<' used to abort the whole PDF export with a parse error.
    """
    return xml_escape(text or "")


def _pdf_inline(line: str) -> str:
    """Escapes a line, then re-applies **bold** as real markup."""
    return re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", _pdf_text(line))


def _split_speaker(block: str):
    """Splits 'Name: text' into ('Name', 'text'), or (None, block) if there is no label.

    Only a short, single-line prefix counts as a speaker label, so a sentence
    that merely contains a colon is not mistaken for one.
    """
    head, sep, tail = block.partition(":")
    if not sep or "\n" in head or len(head.strip()) > 60 or not head.strip():
        return None, block
    return head.strip(), tail.strip()
